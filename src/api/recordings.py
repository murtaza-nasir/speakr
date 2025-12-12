"""
Recording upload, processing, and management.

This blueprint was auto-generated from app.py route extraction.
"""

import os
import json
import re
import mimetypes
import time
import subprocess
from datetime import datetime, timedelta
from src.services.job_queue import job_queue
from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, send_file, Response, current_app, make_response
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge
from sqlalchemy import select
from email.utils import encode_rfc2231

from src.database import db
from src.models import *
from src.utils import *
from src.config.app_config import ASR_MIN_SPEAKERS, ASR_MAX_SPEAKERS, ASR_DIARIZE
from src.tasks.processing import format_transcription_for_llm, transcribe_with_chunking
from src.services.speaker import update_speaker_usage, identify_unidentified_speakers_from_text
from src.services.speaker_embedding_matcher import update_speaker_embedding
from src.services.speaker_snippets import create_speaker_snippets
from src.services.document import process_markdown_to_docx
from src.services.llm import client, chat_client, call_llm_completion, call_chat_completion, process_streaming_with_thinking
from src.services.embeddings import process_recording_chunks
from src.file_exporter import export_recording, mark_export_as_deleted

# Create blueprint
recordings_bp = Blueprint('recordings', __name__)

# Configuration from environment
ENABLE_INQUIRE_MODE = os.environ.get('ENABLE_INQUIRE_MODE', 'false').lower() == 'true'
ENABLE_AUTO_DELETION = os.environ.get('ENABLE_AUTO_DELETION', 'false').lower() == 'true'
DELETION_MODE = os.environ.get('DELETION_MODE', 'full_recording')  # 'audio_only' or 'full_recording'
USERS_CAN_DELETE = os.environ.get('USERS_CAN_DELETE', 'true').lower() == 'true'
ENABLE_INTERNAL_SHARING = os.environ.get('ENABLE_INTERNAL_SHARING', 'false').lower() == 'true'
USE_ASR_ENDPOINT = os.environ.get('USE_ASR_ENDPOINT', 'false').lower() == 'true'
ENABLE_CHUNKING = os.environ.get('ENABLE_CHUNKING', 'true').lower() == 'true'

# Global helpers (will be injected from app)
has_recording_access = None
get_user_recording_status = None
set_user_recording_status = None
enrich_recording_dict_with_user_status = None
bcrypt = None
csrf = None
limiter = None
chunking_service = None

def init_recordings_helpers(**kwargs):
    """Initialize helper functions and extensions from app."""
    global has_recording_access, get_user_recording_status, set_user_recording_status, enrich_recording_dict_with_user_status, bcrypt, csrf, limiter, chunking_service
    has_recording_access = kwargs.get('has_recording_access')
    get_user_recording_status = kwargs.get('get_user_recording_status')
    set_user_recording_status = kwargs.get('set_user_recording_status')
    enrich_recording_dict_with_user_status = kwargs.get('enrich_recording_dict_with_user_status')
    bcrypt = kwargs.get('bcrypt')
    csrf = kwargs.get('csrf')
    limiter = kwargs.get('limiter')
    chunking_service = kwargs.get('chunking_service')


# --- Routes ---

@recordings_bp.route('/recording/<int:recording_id>/download/transcript')
@login_required
def download_transcript_with_template(recording_id):
    """Download transcript with custom template formatting."""
    try:
        import re
        from datetime import timedelta

        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user):
            return jsonify({'error': 'You do not have permission to access this recording'}), 403

        if not recording.transcription:
            return jsonify({'error': 'No transcription available for this recording'}), 400

        # Get template ID from query params
        template_id = request.args.get('template_id', type=int)

        # Get the template
        if template_id:
            template = TranscriptTemplate.query.filter_by(
                id=template_id,
                user_id=current_user.id
            ).first()
        else:
            # Use default template
            template = TranscriptTemplate.query.filter_by(
                user_id=current_user.id,
                is_default=True
            ).first()

        # If no template found, use a basic format
        if not template:
            template_format = "[{{speaker}}]: {{text}}"
        else:
            template_format = template.template

        # Helper functions for formatting
        def format_time(seconds):
            """Format seconds to HH:MM:SS"""
            if seconds is None:
                return "00:00:00"
            td = timedelta(seconds=seconds)
            hours = int(td.total_seconds() // 3600)
            minutes = int((td.total_seconds() % 3600) // 60)
            secs = int(td.total_seconds() % 60)
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"

        def format_srt_time(seconds):
            """Format seconds to SRT format HH:MM:SS,mmm"""
            if seconds is None:
                return "00:00:00,000"
            td = timedelta(seconds=seconds)
            hours = int(td.total_seconds() // 3600)
            minutes = int((td.total_seconds() % 3600) // 60)
            secs = int(td.total_seconds() % 60)
            millis = int((td.total_seconds() % 1) * 1000)
            return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"

        # Parse transcription - handle both JSON (diarized) and plain text formats
        is_diarized = False
        transcription_data = None
        try:
            transcription_data = json.loads(recording.transcription)
            if isinstance(transcription_data, list):
                is_diarized = True
        except (json.JSONDecodeError, TypeError):
            # Not JSON, treat as plain text
            pass

        # If plain text transcription, return it as-is (no template formatting applies)
        if not is_diarized:
            formatted_transcript = recording.transcription
        else:
            # Generate formatted transcript from diarized segments
            output_lines = []
            for index, segment in enumerate(transcription_data, 1):
                line = template_format

                # Replace variables
                replacements = {
                    '{{index}}': str(index),
                    '{{speaker}}': segment.get('speaker', 'Unknown'),
                    '{{text}}': segment.get('sentence', ''),
                    '{{start_time}}': format_time(segment.get('start_time')),
                    '{{end_time}}': format_time(segment.get('end_time')),
                }

                for key, value in replacements.items():
                    line = line.replace(key, value)

                # Handle filters
                # Upper case filter
                line = re.sub(r'{{(.*?)\|upper}}', lambda m: replacements.get('{{' + m.group(1) + '}}', '').upper(), line)
                # SRT time filter
                line = re.sub(r'{{start_time\|srt}}', format_srt_time(segment.get('start_time')), line)
                line = re.sub(r'{{end_time\|srt}}', format_srt_time(segment.get('end_time')), line)

                output_lines.append(line)

            # Join lines
            formatted_transcript = '\n'.join(output_lines)

        # Create response
        response = make_response(formatted_transcript)
        if is_diarized and template:
            filename = f"{recording.title or 'transcript'}_{template.name}.txt"
        elif is_diarized:
            filename = f"{recording.title or 'transcript'}_formatted.txt"
        else:
            # Plain text transcription
            filename = f"{recording.title or 'transcript'}.txt"
        filename = re.sub(r'[^a-zA-Z0-9_\-\.]', '_', filename)
        response.headers['Content-Type'] = 'text/plain; charset=utf-8'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'

        return response

    except Exception as e:
        current_app.logger.error(f"Error downloading transcript: {e}")
        return jsonify({'error': 'Failed to generate transcript download'}), 500




@recordings_bp.route('/recording/<int:recording_id>/download/summary')
@login_required
def download_summary_word(recording_id):
    """Download recording summary as a Word document."""
    try:
        from docx import Document
        from docx.shared import Inches
        import re
        from io import BytesIO
        
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user):
            return jsonify({'error': 'You do not have permission to access this recording'}), 403

        if not recording.summary:
            return jsonify({'error': 'No summary available for this recording'}), 400
        
        # Create Word document
        doc = Document()
        
        # Add title
        title_text = f'Summary: {recording.title or "Untitled Recording"}'
        title = doc.add_heading(title_text, 0)
        # Check if title needs Unicode font support
        try:
            title_text.encode('ascii')
        except UnicodeEncodeError:
            # Title contains non-ASCII characters
            from docx.oxml.ns import qn
            for run in title.runs:
                run.font.name = 'Arial'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        # Helper function to add paragraph with Unicode support
        def add_unicode_paragraph(doc, text):
            p = doc.add_paragraph(text)
            try:
                text.encode('ascii')
            except UnicodeEncodeError:
                from docx.oxml.ns import qn
                for run in p.runs:
                    run.font.name = 'Arial'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            return p

        # Add metadata
        add_unicode_paragraph(doc, f'Uploaded: {recording.created_at.strftime("%Y-%m-%d %H:%M")}')
        if recording.meeting_date:
            add_unicode_paragraph(doc, f'Recording Date: {recording.meeting_date.strftime("%Y-%m-%d")}')
        if recording.participants:
            add_unicode_paragraph(doc, f'Participants: {recording.participants}')
        visible_tags = recording.get_visible_tags(current_user)
        if visible_tags:
            tags_str = ', '.join([tag.name for tag in visible_tags])
            add_unicode_paragraph(doc, f'Tags: {tags_str}')
        doc.add_paragraph('')  # Empty line

        # Process markdown content using the helper function
        process_markdown_to_docx(doc, recording.summary)
        
        # Save to BytesIO
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        
        # Create safe filename
        safe_title = re.sub(r'[<>:"/\\|?*]', '', recording.title or 'Untitled')
        safe_title = re.sub(r'[-\s]+', '-', safe_title).strip('-')
        filename = f'summary-{safe_title}.docx' if safe_title else f'summary-recording-{recording_id}.docx'

        # Create ASCII fallback for send_file - if title has non-ASCII chars, use generic name with ID
        ascii_filename = filename.encode('ascii', 'ignore').decode('ascii')
        if not ascii_filename.strip() or ascii_filename.strip() in ['summary-.docx', 'summary-recording-.docx']:
            ascii_filename = f'summary-recording-{recording_id}.docx'

        response = send_file(
            doc_stream,
            as_attachment=False,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        # Properly encode filename for international characters
        # Check if filename contains non-ASCII characters
        try:
            # Try to encode as ASCII - if this works, use simple format
            filename.encode('ascii')
            # ASCII-only filename, use simple format
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        except UnicodeEncodeError:
            # Contains non-ASCII characters, use proper RFC 2231 encoding
            try:
                # Use Python's built-in RFC 2231 encoder
                encoded_value = encode_rfc2231(filename, charset='utf-8')
                header_value = f'attachment; filename*={encoded_value}'
                current_app.logger.info(f"DEBUG CHINESE FILENAME (RFC2231): Original='{filename}', Header='{header_value}'")
                response.headers['Content-Disposition'] = header_value
            except Exception as e:
                # Fallback to simple attachment with generic name
                current_app.logger.error(f"RFC2231 encoding failed: {e}, using fallback")
                response.headers['Content-Disposition'] = f'attachment; filename="download-{recording_id}.docx"'
        return response
        
    except Exception as e:
        current_app.logger.error(f"Error generating summary Word document: {e}")
        return jsonify({'error': 'Failed to generate Word document'}), 500



@recordings_bp.route('/recording/<int:recording_id>/download/chat', methods=['POST'])
@login_required
def download_chat_word(recording_id):
    """Download chat conversation as a Word document."""
    try:
        from docx import Document
        from docx.shared import Inches
        import re
        from io import BytesIO
        
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user):
            return jsonify({'error': 'You do not have permission to access this recording'}), 403

        # Get chat messages from request
        data = request.json
        if not data or 'messages' not in data:
            return jsonify({'error': 'No messages provided'}), 400
        
        messages = data['messages']
        if not messages:
            return jsonify({'error': 'No messages to download'}), 400
        
        # Create Word document
        doc = Document()
        
        # Add title
        title_text = f'Chat Conversation: {recording.title or "Untitled Recording"}'
        title = doc.add_heading(title_text, 0)
        # Check if title needs Unicode font support
        try:
            title_text.encode('ascii')
        except UnicodeEncodeError:
            from docx.oxml.ns import qn
            for run in title.runs:
                run.font.name = 'Arial'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        
        # Helper function to add paragraph with Unicode support
        def add_unicode_paragraph(doc, text):
            p = doc.add_paragraph(text)
            try:
                text.encode('ascii')
            except UnicodeEncodeError:
                from docx.oxml.ns import qn
                for run in p.runs:
                    run.font.name = 'Arial'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            return p

        # Add metadata
        add_unicode_paragraph(doc, f'Recording Date: {recording.created_at.strftime("%Y-%m-%d %H:%M")}')
        add_unicode_paragraph(doc, f'Chat Export Date: {datetime.utcnow().strftime("%Y-%m-%d %H:%M")}')
        doc.add_paragraph('')  # Empty line
        
        # Add chat messages
        for message in messages:
            role = message.get('role', 'unknown')
            content = message.get('content', '')
            thinking = message.get('thinking', '')
            
            # Add role header
            if role == 'user':
                p = doc.add_paragraph()
                run = p.add_run('You: ')
                run.bold = True
            elif role == 'assistant':
                p = doc.add_paragraph()
                run = p.add_run('Assistant: ')
                run.bold = True
            else:
                p = doc.add_paragraph()
                run = p.add_run(f'{role.title()}: ')
                run.bold = True
            
            # Add thinking content if present
            if thinking and role == 'assistant':
                p = doc.add_paragraph()
                p.add_run('[Model Reasoning]\n').italic = True
                p.add_run(thinking).italic = True
                doc.add_paragraph('')  # Empty line
            
            # Add message content with markdown formatting
            process_markdown_to_docx(doc, content)
            
            doc.add_paragraph('')  # Empty line between messages
        
        # Save to BytesIO
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        
        # Create safe filename
        safe_title = re.sub(r'[<>:"/\\|?*]', '', recording.title or 'Untitled')
        safe_title = re.sub(r'[-\s]+', '-', safe_title).strip('-')
        filename = f'chat-{safe_title}.docx' if safe_title else f'chat-recording-{recording_id}.docx'

        # Create ASCII fallback for send_file - if title has non-ASCII chars, use generic name with ID
        ascii_filename = filename.encode('ascii', 'ignore').decode('ascii')
        if not ascii_filename.strip() or ascii_filename.strip() in ['chat-.docx', 'chat-recording-.docx']:
            ascii_filename = f'chat-recording-{recording_id}.docx'

        response = send_file(
            doc_stream,
            as_attachment=False,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        # Properly encode filename for international characters
        # Check if filename contains non-ASCII characters
        try:
            # Try to encode as ASCII - if this works, use simple format
            filename.encode('ascii')
            # ASCII-only filename, use simple format
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        except UnicodeEncodeError:
            # Contains non-ASCII characters, use proper RFC 2231 encoding
            try:
                # Use Python's built-in RFC 2231 encoder
                encoded_value = encode_rfc2231(filename, charset='utf-8')
                header_value = f'attachment; filename*={encoded_value}'
                current_app.logger.info(f"DEBUG CHINESE FILENAME (RFC2231): Original='{filename}', Header='{header_value}'")
                response.headers['Content-Disposition'] = header_value
            except Exception as e:
                # Fallback to simple attachment with generic name
                current_app.logger.error(f"RFC2231 encoding failed: {e}, using fallback")
                response.headers['Content-Disposition'] = f'attachment; filename="download-{recording_id}.docx"'
        return response
        
    except Exception as e:
        current_app.logger.error(f"Error generating chat Word document: {e}")
        return jsonify({'error': 'Failed to generate Word document'}), 500



@recordings_bp.route('/recording/<int:recording_id>/download/notes')
@login_required
def download_notes_word(recording_id):
    """Download recording notes as a Word document."""
    try:
        from docx import Document
        from docx.shared import Inches
        import re
        from io import BytesIO
        
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user):
            return jsonify({'error': 'You do not have permission to access this recording'}), 403

        if not recording.notes:
            return jsonify({'error': 'No notes available for this recording'}), 400
        
        # Create Word document
        doc = Document()
        
        # Add title
        title_text = f'Notes: {recording.title or "Untitled Recording"}'
        title = doc.add_heading(title_text, 0)
        # Check if title needs Unicode font support
        try:
            title_text.encode('ascii')
        except UnicodeEncodeError:
            from docx.oxml.ns import qn
            for run in title.runs:
                run.font.name = 'Arial'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        # Helper function to add paragraph with Unicode support
        def add_unicode_paragraph(doc, text):
            p = doc.add_paragraph(text)
            try:
                text.encode('ascii')
            except UnicodeEncodeError:
                from docx.oxml.ns import qn
                for run in p.runs:
                    run.font.name = 'Arial'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            return p

        # Add metadata
        add_unicode_paragraph(doc, f'Uploaded: {recording.created_at.strftime("%Y-%m-%d %H:%M")}')
        if recording.meeting_date:
            add_unicode_paragraph(doc, f'Recording Date: {recording.meeting_date.strftime("%Y-%m-%d")}')
        if recording.participants:
            add_unicode_paragraph(doc, f'Participants: {recording.participants}')
        visible_tags = recording.get_visible_tags(current_user)
        if visible_tags:
            tags_str = ', '.join([tag.name for tag in visible_tags])
            add_unicode_paragraph(doc, f'Tags: {tags_str}')
        doc.add_paragraph('')  # Empty line

        # Process markdown content using the helper function
        process_markdown_to_docx(doc, recording.notes)
        
        # Save to BytesIO
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        
        # Create safe filename
        safe_title = re.sub(r'[<>:"/\\|?*]', '', recording.title or 'Untitled')
        safe_title = re.sub(r'[-\s]+', '-', safe_title).strip('-')
        filename = f'notes-{safe_title}.docx' if safe_title else f'notes-recording-{recording_id}.docx'

        # Create ASCII fallback for send_file - if title has non-ASCII chars, use generic name with ID
        ascii_filename = filename.encode('ascii', 'ignore').decode('ascii')
        if not ascii_filename.strip() or ascii_filename.strip() in ['notes-.docx', 'notes-recording-.docx']:
            ascii_filename = f'notes-recording-{recording_id}.docx'

        response = send_file(
            doc_stream,
            as_attachment=False,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        # Properly encode filename for international characters
        # Check if filename contains non-ASCII characters
        try:
            # Try to encode as ASCII - if this works, use simple format
            filename.encode('ascii')
            # ASCII-only filename, use simple format
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        except UnicodeEncodeError:
            # Contains non-ASCII characters, use proper RFC 2231 encoding
            try:
                # Use Python's built-in RFC 2231 encoder
                encoded_value = encode_rfc2231(filename, charset='utf-8')
                header_value = f'attachment; filename*={encoded_value}'
                current_app.logger.info(f"DEBUG CHINESE FILENAME (RFC2231): Original='{filename}', Header='{header_value}'")
                response.headers['Content-Disposition'] = header_value
            except Exception as e:
                # Fallback to simple attachment with generic name
                current_app.logger.error(f"RFC2231 encoding failed: {e}, using fallback")
                response.headers['Content-Disposition'] = f'attachment; filename="download-{recording_id}.docx"'
        return response
        
    except Exception as e:
        current_app.logger.error(f"Error generating notes Word document: {e}")
        return jsonify({'error': 'Failed to generate Word document'}), 500



@recordings_bp.route('/recording/<int:recording_id>/generate_summary', methods=['POST'])
@login_required
def generate_summary_endpoint(recording_id):
    """Generate summary for a recording that doesn't have one."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to generate summary for this recording'}), 403
            
        # Check if transcription exists
        if not recording.transcription or len(recording.transcription.strip()) < 10:
            return jsonify({'error': 'No valid transcription available for summary generation'}), 400
            
        # Check if already processing
        if recording.status in ['PROCESSING', 'SUMMARIZING']:
            return jsonify({'error': 'Recording is already being processed'}), 400
            
        # Check if OpenRouter client is available
        if client is None:
            return jsonify({'error': 'Summary service is not available (OpenRouter client not configured)'}), 503
            
        current_app.logger.info(f"Queueing summary generation for recording {recording_id}")

        # Queue summary generation job
        job_queue.enqueue(
            user_id=current_user.id,
            recording_id=recording.id,
            job_type='summarize',
            params={'user_id': current_user.id}
        )

        return jsonify({
            'success': True,
            'message': 'Summary generation queued'
        })
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error starting summary generation for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}), 500



@recordings_bp.route('/recording/<int:recording_id>/update_speakers', methods=['POST'])
@login_required
def update_speakers(recording_id):
    """Updates speaker labels in a transcription with provided names."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to edit this recording'}), 403

        data = request.json
        speaker_map = data.get('speaker_map')
        regenerate_summary = data.get('regenerate_summary', False)

        if speaker_map is None:
            return jsonify({'error': 'No speaker map provided'}), 400

        transcription_text = recording.transcription
        is_json = False
        try:
            transcription_data = json.loads(transcription_text)
            # Updated check for our new simplified JSON format (a list of segment objects)
            is_json = isinstance(transcription_data, list)
        except (json.JSONDecodeError, TypeError):
            is_json = False

        speaker_names_used = []

        if is_json:
            # Handle new simplified JSON transcript (list of segments)
            for segment in transcription_data:
                original_speaker_label = segment.get('speaker')
                if original_speaker_label in speaker_map:
                    new_name_info = speaker_map[original_speaker_label]
                    new_name = new_name_info.get('name', '').strip()
                    # If isMe is checked but no name provided, use current user's name
                    if new_name_info.get('isMe') and not new_name:
                        new_name = current_user.name or 'Me'

                    if new_name:
                        segment['speaker'] = new_name
                        if new_name not in speaker_names_used:
                            speaker_names_used.append(new_name)
            
            recording.transcription = json.dumps(transcription_data)
            
            # Update participants only from speakers that were actually given names (not default labels)
            final_speakers = set()
            for seg in transcription_data:
                speaker = seg.get('speaker')
                if speaker and str(speaker).strip():
                    # Only include speakers that have been given actual names (not default labels like "SPEAKER_01", "SPEAKER_09", etc.)
                    # Check if this speaker was updated with a real name (not a default label)
                    if not re.match(r'^SPEAKER_\d+$', str(speaker), re.IGNORECASE):
                        final_speakers.add(speaker)
            recording.participants = ', '.join(sorted(list(final_speakers)))

        else:
            # Handle plain text transcript
            new_participants = []
            for speaker_label, new_name_info in speaker_map.items():
                new_name = new_name_info.get('name', '').strip()
                # If isMe is checked but no name provided, use current user's name
                if new_name_info.get('isMe') and not new_name:
                    new_name = current_user.name or 'Me'

                if new_name:
                    transcription_text = re.sub(r'\[\s*' + re.escape(speaker_label) + r'\s*\]', f'[{new_name}]', transcription_text, flags=re.IGNORECASE)
                    if new_name not in new_participants:
                        new_participants.append(new_name)
            
            recording.transcription = transcription_text
            if new_participants:
                recording.participants = ', '.join(new_participants)
            speaker_names_used = new_participants

        # Update speaker usage statistics
        if speaker_names_used:
            update_speaker_usage(speaker_names_used)

        # Update speaker voice embeddings if available
        embeddings_updated = 0
        snippets_created = 0
        if recording.speaker_embeddings and speaker_map:
            try:
                # Parse embeddings from recording
                embeddings_data = json.loads(recording.speaker_embeddings) if isinstance(recording.speaker_embeddings, str) else recording.speaker_embeddings

                # Build reverse map: SPEAKER_XX -> actual name assigned
                speaker_label_to_name = {}
                for speaker_label, speaker_info in speaker_map.items():
                    name = speaker_info.get('name', '').strip()
                    # Handle isMe checkbox
                    if speaker_info.get('isMe') and not name:
                        name = current_user.name or 'Me'

                    # Only include speakers that were given real names (not SPEAKER_XX)
                    if name and not re.match(r'^SPEAKER_\d+$', name, re.IGNORECASE):
                        speaker_label_to_name[speaker_label] = name

                # Update embeddings for each identified speaker
                for speaker_label, embedding in embeddings_data.items():
                    if speaker_label in speaker_label_to_name and embedding and len(embedding) == 256:
                        speaker_name = speaker_label_to_name[speaker_label]

                        # Find or create the speaker
                        speaker = Speaker.query.filter_by(
                            user_id=current_user.id,
                            name=speaker_name
                        ).first()

                        if speaker:
                            # Update the speaker's voice embedding
                            similarity = update_speaker_embedding(speaker, embedding, recording.id)
                            embeddings_updated += 1

                            if similarity is not None:
                                current_app.logger.info(
                                    f"Updated voice profile for '{speaker_name}' "
                                    f"(similarity: {similarity*100:.1f}%)"
                                )
                            else:
                                current_app.logger.info(
                                    f"Created initial voice profile for '{speaker_name}'"
                                )

                # Create snippets for identified speakers
                if speaker_label_to_name:
                    snippets_created = create_speaker_snippets(recording.id, speaker_map)
                    if snippets_created > 0:
                        current_app.logger.info(f"Created {snippets_created} speaker snippets")

            except Exception as e:
                current_app.logger.error(f"Error updating speaker embeddings: {e}", exc_info=True)
                # Don't fail the whole request if embedding update fails

        db.session.commit()

        summary_queued = False
        if regenerate_summary:
            current_app.logger.info(f"Queueing summary regeneration for recording {recording_id} after speaker update.")
            job_queue.enqueue(
                user_id=current_user.id,
                recording_id=recording.id,
                job_type='summarize',
                params={'user_id': current_user.id}
            )
            summary_queued = True

        # Return recording with per-user status
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({
            'success': True,
            'message': 'Speakers updated successfully.',
            'recording': recording_dict,
            'summary_queued': summary_queued
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error updating speakers for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500



@recordings_bp.route('/recording/<int:recording_id>/update_transcript', methods=['POST'])
@login_required
def update_transcript(recording_id):
    """Updates the complete transcript data including text edits and speaker changes."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to edit this recording'}), 403

        data = request.json
        transcript_data = data.get('transcript_data')
        speaker_map = data.get('speaker_map', {})
        regenerate_summary = data.get('regenerate_summary', False)

        if not transcript_data or not isinstance(transcript_data, list):
            return jsonify({'error': 'Invalid transcript data provided'}), 400

        # Update speaker names in the transcript data
        speaker_names_used = []
        for segment in transcript_data:
            original_speaker_label = segment.get('speaker')

            # Apply speaker name mapping if provided
            if original_speaker_label in speaker_map:
                new_name_info = speaker_map[original_speaker_label]
                new_name = new_name_info.get('name', '').strip()
                if new_name_info.get('isMe'):
                    new_name = current_user.name or 'Me'

                if new_name:
                    segment['speaker'] = new_name
                    if new_name not in speaker_names_used:
                        speaker_names_used.append(new_name)

        # Save the updated transcript
        recording.transcription = json.dumps(transcript_data)

        # Update participants
        final_speakers = set()
        for seg in transcript_data:
            speaker = seg.get('speaker')
            if speaker and str(speaker).strip():
                # Only include speakers with real names (not default labels)
                if not re.match(r'^SPEAKER_\d+$', str(speaker), re.IGNORECASE):
                    final_speakers.add(speaker)
        recording.participants = ', '.join(sorted(list(final_speakers)))

        # Update speaker usage statistics
        if speaker_names_used:
            update_speaker_usage(speaker_names_used)

        db.session.commit()

        summary_queued = False
        if regenerate_summary:
            current_app.logger.info(f"Queueing summary regeneration for recording {recording_id} after transcript update.")
            job_queue.enqueue(
                user_id=current_user.id,
                recording_id=recording.id,
                job_type='summarize',
                params={'user_id': current_user.id}
            )
            summary_queued = True
            # Export will happen after summary regenerates
        else:
            # Re-export the recording if auto-export is enabled
            export_recording(recording_id)

        # Return recording with per-user status
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({
            'success': True,
            'message': 'Transcript updated successfully.',
            'recording': recording_dict,
            'summary_queued': summary_queued
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error updating transcript for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500



@recordings_bp.route('/recording/<int:recording_id>/auto_identify_speakers', methods=['POST'])
@login_required
def auto_identify_speakers(recording_id):
    """
    Automatically identifies speakers in a transcription using an LLM.
    Strips existing names and re-identifies all speakers from scratch.
    """
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user):
            return jsonify({'error': 'You do not have permission to modify this recording'}), 403

        if not recording.transcription:
            return jsonify({'error': 'No transcription available for speaker identification'}), 400

        # Parse the transcription
        try:
            transcription_data = json.loads(recording.transcription)
        except (json.JSONDecodeError, TypeError):
            return jsonify({'error': 'Transcription format not supported for auto-identification'}), 400

        if not isinstance(transcription_data, list):
            return jsonify({'error': 'Transcription format not supported for auto-identification'}), 400

        # Extract unique speakers in order of appearance
        seen_speakers = set()
        unique_speakers = []
        for segment in transcription_data:
            speaker = segment.get('speaker')
            if speaker and speaker not in seen_speakers:
                seen_speakers.add(speaker)
                unique_speakers.append(speaker)

        if not unique_speakers:
            return jsonify({'error': 'No speakers found in transcription'}), 400

        # Create a mapping from current names to SPEAKER_XX labels
        speaker_to_label = {}
        for idx, speaker in enumerate(unique_speakers):
            speaker_to_label[speaker] = f'SPEAKER_{str(idx).zfill(2)}'

        # Create a temporary transcript with SPEAKER_XX labels
        temp_transcript = []
        for segment in transcription_data:
            temp_segment = segment.copy()
            original_speaker = segment.get('speaker')
            if original_speaker:
                temp_segment['speaker'] = speaker_to_label[original_speaker]
            temp_transcript.append(temp_segment)

        # Format for LLM directly (don't pass JSON, pass formatted text)
        formatted_lines = []
        for segment in temp_transcript:
            speaker = segment.get('speaker', 'Unknown Speaker')
            sentence = segment.get('sentence', '')
            formatted_lines.append(f"[{speaker}]: {sentence}")
        formatted_transcription = "\n".join(formatted_lines)

        # Get all SPEAKER_XX labels
        speaker_labels = list(speaker_to_label.values())

        current_app.logger.info(f"[Auto-Identify] Formatted transcript (first 500 chars): {formatted_transcription[:500]}")
        current_app.logger.info(f"[Auto-Identify] Speaker labels: {speaker_labels}")

        # Call identify_unidentified_speakers_from_text but pass the formatted text directly
        # We need to bypass format_transcription_for_llm since we already formatted it
        if not speaker_labels:
            return jsonify({'error': 'No speakers found in transcription'}), 400

        # Get configurable transcript length limit
        transcript_limit = SystemSetting.get_setting('transcript_length_limit', 30000)
        if transcript_limit == -1:
            transcript_text = formatted_transcription
        else:
            transcript_text = formatted_transcription[:transcript_limit]

        prompt = f"""Analyze the following conversation transcript and identify the names of the speakers based on the context and content of their dialogue.

The speakers that need to be identified are: {', '.join(speaker_labels)}

Look for clues in the conversation such as:
- Names mentioned by other speakers when addressing someone
- Self-introductions or references to their own name
- Context clues about roles, relationships, or positions
- Any direct mentions of names in the dialogue

Here is the complete conversation transcript:

{transcript_text}

Based on the conversation above, identify the most likely real names for the speakers. Pay close attention to how speakers address each other and any names that are mentioned in the dialogue.

Respond with a single JSON object where keys are the speaker labels (e.g., "SPEAKER_01") and values are the identified full names. If a name cannot be determined from the conversation context, use an empty string "".

Example format:
{{
  "SPEAKER_01": "Jane Smith",
  "SPEAKER_03": "Bob Johnson",
  "SPEAKER_05": ""
}}

JSON Response:
"""

        try:
            from src.services.llm import call_llm_completion
            current_app.logger.info(f"[Auto-Identify] Calling LLM directly with prompt")

            completion = call_llm_completion(
                messages=[
                    {"role": "system", "content": "You are an expert in analyzing conversation transcripts to identify speakers based on contextual clues in the dialogue. Analyze the conversation carefully to find names mentioned when speakers address each other or introduce themselves. Your response must be a single, valid JSON object containing only the requested speaker identifications."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2,
                response_format={"type": "json_object"}
            )
            response_content = completion.choices[0].message.content
            current_app.logger.info(f"[Auto-Identify] LLM Raw Response: {response_content}")

            from src.utils import safe_json_loads
            identified_map = safe_json_loads(response_content, {})
            current_app.logger.info(f"[Auto-Identify] Parsed identified_map: {identified_map}")

            # Post-process the map to replace "Unknown" with an empty string
            for speaker_label, identified_name in identified_map.items():
                if identified_name and identified_name.strip().lower() in ["unknown", "n/a", "not available", "unclear"]:
                    identified_map[speaker_label] = ""
        except Exception as e:
            current_app.logger.error(f"[Auto-Identify] Error calling LLM: {e}", exc_info=True)
            return jsonify({'error': f'Failed to identify speakers: {str(e)}'}), 500
        current_app.logger.info(f"[Auto-Identify] LLM returned identified_map: {identified_map}")

        # Map back to original speaker IDs
        final_speaker_map = {}
        for original_speaker, temp_label in speaker_to_label.items():
            if temp_label in identified_map:
                final_speaker_map[original_speaker] = identified_map[temp_label]

        current_app.logger.info(f"[Auto-Identify] Final speaker_map to return: {final_speaker_map}")
        current_app.logger.info(f"[Auto-Identify] Speaker mapping: {speaker_to_label}")

        return jsonify({'success': True, 'speaker_map': final_speaker_map})

    except ValueError as ve:
        # Handle cases where API key is not set
        return jsonify({'error': str(ve)}), 503
    except Exception as e:
        current_app.logger.error(f"Error during auto speaker identification for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': f'An unexpected error occurred: {str(e)}'}), 500

# --- Chat with Transcription ---


@recordings_bp.route('/recording/<int:recording_id>/reprocess_transcription', methods=['POST'])
@login_required
def reprocess_transcription(recording_id):
    """Reprocess transcription for a given recording."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to reprocess this recording'}), 403

        if not recording.audio_path or not os.path.exists(recording.audio_path):
            return jsonify({'error': 'Audio file not found for reprocessing'}), 404

        if recording.status in ['QUEUED', 'PROCESSING', 'SUMMARIZING']:
            return jsonify({'error': 'Recording is already being processed'}), 400

        # File path and name for processing (conversion handled in background task if needed)
        filepath = recording.audio_path
        filename_for_asr = recording.original_filename or os.path.basename(filepath)

        # --- Proceed with reprocessing ---
        recording.transcription = None
        recording.summary = None
        recording.status = 'QUEUED'  # Will change to PROCESSING when job starts

        # Clear existing events since they depend on the transcription
        Event.query.filter_by(recording_id=recording_id).delete()

        db.session.commit()

        current_app.logger.info(f"Queueing transcription reprocessing for recording {recording_id}")

        # Prepare job parameters
        data = request.json or {}
        start_time = datetime.utcnow()
        app_context = current_app._get_current_object().app_context()

        # Decide which transcription method to use
        if USE_ASR_ENDPOINT:
            language = data.get('language') or (recording.owner.transcription_language if recording.owner else None)
            min_speakers = data.get('min_speakers') or None
            max_speakers = data.get('max_speakers') or None

            # Convert to int if provided
            if min_speakers:
                try:
                    min_speakers = int(min_speakers)
                except (ValueError, TypeError):
                    min_speakers = None
            if max_speakers:
                try:
                    max_speakers = int(max_speakers)
                except (ValueError, TypeError):
                    max_speakers = None

            # Apply tag defaults if no user input provided
            if (min_speakers is None or max_speakers is None) and recording.tags:
                for tag_association in sorted(recording.tag_associations, key=lambda x: x.order):
                    tag = tag_association.tag
                    if min_speakers is None and tag.default_min_speakers:
                        min_speakers = tag.default_min_speakers
                    if max_speakers is None and tag.default_max_speakers:
                        max_speakers = tag.default_max_speakers
                    if min_speakers is not None and max_speakers is not None:
                        break

            # Apply environment variable defaults
            if min_speakers is None and ASR_MIN_SPEAKERS:
                try:
                    min_speakers = int(ASR_MIN_SPEAKERS)
                except (ValueError, TypeError):
                    min_speakers = None
            if max_speakers is None and ASR_MAX_SPEAKERS:
                try:
                    max_speakers = int(ASR_MAX_SPEAKERS)
                except (ValueError, TypeError):
                    max_speakers = None

            # Enqueue the job using new fair queue API
            job_params = {
                'language': language,
                'min_speakers': min_speakers,
                'max_speakers': max_speakers
            }
        else:
            # Standard Whisper API - no special params needed
            job_params = {}

        job_id = job_queue.enqueue(
            user_id=current_user.id,
            recording_id=recording.id,
            job_type='reprocess_transcription',
            params=job_params
        )

        # Get queue position for response
        queue_position = job_queue.get_position_in_queue(recording.id)
        queue_status = job_queue.get_queue_status()

        # Return recording with per-user status and queue info
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({
            'success': True,
            'message': 'Transcription reprocessing queued',
            'recording': recording_dict,
            'queue_position': queue_position,
            'queue_status': queue_status
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error reprocessing transcription for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}), 500




@recordings_bp.route('/recording/<int:recording_id>/reprocess_summary', methods=['POST'])
@login_required
def reprocess_summary(recording_id):
    """Reprocess summary for a given recording (requires existing transcription)."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to reprocess this recording'}), 403
            
        # Check if transcription exists
        if not recording.transcription or len(recording.transcription.strip()) < 10:
            return jsonify({'error': 'No valid transcription available for summary generation'}), 400
            
        # Check if already processing
        if recording.status in ['PROCESSING', 'SUMMARIZING']:
            return jsonify({'error': 'Recording is already being processed'}), 400
            
        # Check if OpenRouter client is available
        if client is None:
            return jsonify({'error': 'Summary service is not available (OpenRouter client not configured)'}), 503

        # Get custom prompt from request if provided
        data = request.get_json() or {}
        custom_prompt = data.get('custom_prompt', '').strip() if data.get('custom_prompt') else None

        # Debug logging
        if custom_prompt:
            current_app.logger.info(f"Received custom prompt override for recording {recording_id} (length: {len(custom_prompt)})")
        else:
            current_app.logger.info(f"No custom prompt override provided for recording {recording_id}, will use default priority")

        # Clear existing summary (status will be set to QUEUED by job_queue.enqueue)
        recording.summary = None

        # Clear existing events since they might be re-extracted during summary generation
        Event.query.filter_by(recording_id=recording_id).delete()

        db.session.commit()

        current_app.logger.info(f"Queueing summary reprocessing for recording {recording_id}" +
                       (f" with custom prompt (length: {len(custom_prompt)})" if custom_prompt else ""))

        # Queue summary generation job
        job_params = {
            'custom_prompt': custom_prompt,
            'user_id': current_user.id
        }
        job_queue.enqueue(
            user_id=current_user.id,
            recording_id=recording.id,
            job_type='reprocess_summary',
            params=job_params
        )

        # Refresh recording to get updated status
        db.session.refresh(recording)

        # Return recording with per-user status
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({
            'success': True,
            'message': 'Summary reprocessing started',
            'recording': recording_dict
        })
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error reprocessing summary for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}), 500



@recordings_bp.route('/recording/<int:recording_id>/reset_status', methods=['POST'])
@login_required
def reset_status(recording_id):
    """Resets the status of a stuck or failed recording."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to modify this recording'}), 403

        # Allow resetting if it's stuck or failed
        if recording.status in ['PENDING', 'PROCESSING', 'SUMMARIZING', 'FAILED']:
            recording.status = 'FAILED'
            recording.error_message = "Manually reset from stuck or failed state."
            db.session.commit()
            current_app.logger.info(f"Manually reset status for recording {recording_id} to FAILED.")

            # Return recording with per-user status
            recording_dict = recording.to_dict(viewer_user=current_user)
            enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
            return jsonify({'success': True, 'message': 'Recording status has been reset.', 'recording': recording_dict})
        else:
            return jsonify({'error': f'Recording is not in a state that can be reset. Current status: {recording.status}'}), 400

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error resetting status for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}), 500

# --- Authentication Routes ---


@recordings_bp.route('/')
@login_required
def index():
    # Check if user is a group admin
    is_team_admin = GroupMembership.query.filter_by(
        user_id=current_user.id,
        role='admin'
    ).first() is not None

    # Pass the ASR config, inquire mode config, and user language preference to the template
    user_language = current_user.ui_language if current_user.is_authenticated and current_user.ui_language else 'en'

    # Calculate if archive toggle should be shown (only when audio-only deletion mode is active)
    enable_archive_toggle = ENABLE_AUTO_DELETION and DELETION_MODE == 'audio_only'

    return render_template('index.html',
                         use_asr_endpoint=USE_ASR_ENDPOINT,
                         inquire_mode_enabled=ENABLE_INQUIRE_MODE,
                         enable_archive_toggle=enable_archive_toggle,
                         enable_internal_sharing=ENABLE_INTERNAL_SHARING,
                         user_language=user_language,
                         is_team_admin=is_team_admin)



def get_accessible_recording_ids(user_id):
    """
    Get all recording IDs that a user has access to.

    Includes:
    - Recordings owned by the user
    - Recordings shared with the user via InternalShare
    - Recordings shared via group tags (if team membership exists)

    Args:
        user_id (int): User ID to check access for

    Returns:
        list: List of recording IDs the user can access
    """
    accessible_ids = set()

    # 1. User's own recordings
    own_recordings = db.session.query(Recording.id).filter_by(user_id=user_id).all()
    accessible_ids.update([r.id for r in own_recordings])

    # 2. Internally shared recordings
    if ENABLE_INTERNAL_SHARING:
        shared_recordings = db.session.query(InternalShare.recording_id).filter_by(
            shared_with_user_id=user_id
        ).all()
        accessible_ids.update([r.recording_id for r in shared_recordings])

    return list(accessible_ids)


@recordings_bp.route('/recordings', methods=['GET'])
def get_recordings():
    """Get all recordings for the current user (simple list)."""
    try:
        # Check if user is logged in
        if not current_user.is_authenticated:
            return jsonify([])  # Return empty array if not logged in

        # Filter recordings by the current user
        stmt = select(Recording).where(Recording.user_id == current_user.id).order_by(Recording.created_at.desc())
        recordings = db.session.execute(stmt).scalars().all()
        return jsonify([recording.to_dict(viewer_user=current_user) for recording in recordings])
    except Exception as e:
        current_app.logger.error(f"Error fetching recordings: {e}")
        return jsonify({'error': str(e)}), 500


@recordings_bp.route('/api/recordings', methods=['GET'])
@login_required
def get_recordings_paginated():
    """Get recordings with pagination and server-side filtering (includes shared recordings)."""
    import re
    try:
        # Parse query parameters
        page = request.args.get('page', 1, type=int)
        per_page = min(request.args.get('per_page', 25, type=int), 100)  # Cap at 100 per page
        search_query = request.args.get('q', '').strip()
        show_archived = request.args.get('archived', '').lower() == 'true'
        show_shared = request.args.get('shared', '').lower() == 'true'

        # Get all accessible recording IDs (own + shared)
        accessible_recording_ids = get_accessible_recording_ids(current_user.id)

        if not accessible_recording_ids:
            return jsonify({
                'recordings': [],
                'pagination': {
                    'page': page,
                    'per_page': per_page,
                    'total': 0,
                    'total_pages': 0,
                    'has_next': False,
                    'has_prev': False
                }
            })

        # Build base query to include accessible recordings
        stmt = select(Recording).where(Recording.id.in_(accessible_recording_ids))

        # Apply archived filter (AND with other filters)
        if show_archived:
            # Only show recordings where audio has been deleted
            stmt = stmt.where(Recording.audio_deleted_at.is_not(None))

        # Apply shared filter (AND with other filters)
        if show_shared:
            # Only show recordings shared with current user (not owned by them)
            stmt = stmt.where(Recording.user_id != current_user.id)

        # Apply search filters if provided
        if search_query:
            # Extract date filters
            date_filters = re.findall(r'date:(\S+)', search_query.lower())
            date_from_filters = re.findall(r'date_from:(\S+)', search_query.lower())
            date_to_filters = re.findall(r'date_to:(\S+)', search_query.lower())
            tag_filters = re.findall(r'tag:(\S+)', search_query.lower())
            speaker_filters = re.findall(r'speaker:(\S+)', search_query.lower())

            # Remove special syntax to get text search
            text_query = re.sub(r'date:\S+', '', search_query, flags=re.IGNORECASE)
            text_query = re.sub(r'date_from:\S+', '', text_query, flags=re.IGNORECASE)
            text_query = re.sub(r'date_to:\S+', '', text_query, flags=re.IGNORECASE)
            text_query = re.sub(r'tag:\S+', '', text_query, flags=re.IGNORECASE)
            text_query = re.sub(r'speaker:\S+', '', text_query, flags=re.IGNORECASE).strip()

            # Apply date filters
            for date_filter in date_filters:
                if date_filter == 'today':
                    today = datetime.now().date()
                    stmt = stmt.where(
                        db.or_(
                            db.func.date(Recording.meeting_date) == today,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) == today
                            )
                        )
                    )
                elif date_filter == 'yesterday':
                    yesterday = datetime.now().date() - timedelta(days=1)
                    stmt = stmt.where(
                        db.or_(
                            db.func.date(Recording.meeting_date) == yesterday,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) == yesterday
                            )
                        )
                    )
                elif date_filter == 'thisweek':
                    today = datetime.now().date()
                    start_of_week = today - timedelta(days=today.weekday())
                    stmt = stmt.where(
                        db.or_(
                            Recording.meeting_date >= start_of_week,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= start_of_week
                            )
                        )
                    )
                elif date_filter == 'lastweek':
                    today = datetime.now().date()
                    end_of_last_week = today - timedelta(days=today.weekday())
                    start_of_last_week = end_of_last_week - timedelta(days=7)
                    stmt = stmt.where(
                        db.or_(
                            db.and_(
                                Recording.meeting_date >= start_of_last_week,
                                Recording.meeting_date < end_of_last_week
                            ),
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= start_of_last_week,
                                db.func.date(Recording.created_at) < end_of_last_week
                            )
                        )
                    )
                elif date_filter == 'thismonth':
                    today = datetime.now().date()
                    start_of_month = today.replace(day=1)
                    stmt = stmt.where(
                        db.or_(
                            Recording.meeting_date >= start_of_month,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= start_of_month
                            )
                        )
                    )
                elif date_filter == 'lastmonth':
                    today = datetime.now().date()
                    first_day_this_month = today.replace(day=1)
                    last_day_last_month = first_day_this_month - timedelta(days=1)
                    first_day_last_month = last_day_last_month.replace(day=1)
                    stmt = stmt.where(
                        db.or_(
                            db.and_(
                                Recording.meeting_date >= first_day_last_month,
                                Recording.meeting_date <= last_day_last_month
                            ),
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= first_day_last_month,
                                db.func.date(Recording.created_at) <= last_day_last_month
                            )
                        )
                    )
                elif re.match(r'^\d{4}-\d{2}-\d{2}$', date_filter):
                    # Specific date format YYYY-MM-DD
                    target_date = datetime.strptime(date_filter, '%Y-%m-%d').date()
                    stmt = stmt.where(
                        db.or_(
                            db.func.date(Recording.meeting_date) == target_date,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) == target_date
                            )
                        )
                    )
                elif re.match(r'^\d{4}-\d{2}$', date_filter):
                    # Month format YYYY-MM
                    year, month = map(int, date_filter.split('-'))
                    stmt = stmt.where(
                        db.or_(
                            db.and_(
                                db.extract('year', Recording.meeting_date) == year,
                                db.extract('month', Recording.meeting_date) == month
                            ),
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.extract('year', Recording.created_at) == year,
                                db.extract('month', Recording.created_at) == month
                            )
                        )
                    )
                elif re.match(r'^\d{4}$', date_filter):
                    # Year format YYYY
                    year = int(date_filter)
                    stmt = stmt.where(
                        db.or_(
                            db.extract('year', Recording.meeting_date) == year,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.extract('year', Recording.created_at) == year
                            )
                        )
                    )

            # Apply date range filters
            if date_from_filters and date_from_filters[0]:
                try:
                    date_from = datetime.strptime(date_from_filters[0], '%Y-%m-%d').date()
                    stmt = stmt.where(
                        db.or_(
                            Recording.meeting_date >= date_from,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) >= date_from
                            )
                        )
                    )
                except ValueError:
                    pass  # Invalid date format, ignore

            if date_to_filters and date_to_filters[0]:
                try:
                    date_to = datetime.strptime(date_to_filters[0], '%Y-%m-%d').date()
                    stmt = stmt.where(
                        db.or_(
                            Recording.meeting_date <= date_to,
                            db.and_(
                                Recording.meeting_date.is_(None),
                                db.func.date(Recording.created_at) <= date_to
                            )
                        )
                    )
                except ValueError:
                    pass  # Invalid date format, ignore

            # Apply tag filters
            if tag_filters:
                # Join with tags table and filter by tag names
                tag_conditions = []
                for tag_filter in tag_filters:
                    # Replace underscores back to spaces for matching
                    tag_name = tag_filter.replace('_', ' ')
                    tag_conditions.append(Tag.name.ilike(f'%{tag_name}%'))

                stmt = stmt.join(RecordingTag).join(Tag).where(db.or_(*tag_conditions))

            # Apply speaker filters
            if speaker_filters:
                speaker_conditions = []
                for speaker_filter in speaker_filters:
                    # Replace underscores back to spaces for matching
                    speaker_name = speaker_filter.replace('_', ' ')
                    speaker_conditions.append(Recording.participants.ilike(f'%{speaker_name}%'))
                stmt = stmt.where(db.or_(*speaker_conditions))

            # Apply text search
            if text_query:
                from src.models.sharing import SharedRecordingState

                # Search in user-specific notes:
                # - For owned recordings: search Recording.notes
                # - For shared recordings: search SharedRecordingState.personal_notes

                text_conditions = [
                    Recording.title.ilike(f'%{text_query}%'),
                    Recording.participants.ilike(f'%{text_query}%'),
                    Recording.transcription.ilike(f'%{text_query}%'),
                    # Search owner's notes for owned recordings
                    db.and_(
                        Recording.user_id == current_user.id,
                        Recording.notes.ilike(f'%{text_query}%')
                    )
                ]

                # Add search for personal notes in shared recordings
                # Use a subquery to check if personal_notes match
                shared_notes_subq = select(SharedRecordingState.recording_id).where(
                    db.and_(
                        SharedRecordingState.user_id == current_user.id,
                        SharedRecordingState.personal_notes.ilike(f'%{text_query}%')
                    )
                ).scalar_subquery()

                text_conditions.append(Recording.id.in_(shared_notes_subq))

                stmt = stmt.where(db.or_(*text_conditions))

        # Apply ordering (most recent first based on meeting_date or created_at)
        stmt = stmt.order_by(
            db.case(
                (Recording.meeting_date.is_not(None), Recording.meeting_date),
                else_=db.func.date(Recording.created_at)
            ).desc(),
            Recording.created_at.desc()
        )

        # Get total count for pagination info
        count_stmt = select(db.func.count()).select_from(stmt.subquery())
        total_count = db.session.execute(count_stmt).scalar()

        # Apply pagination
        offset = (page - 1) * per_page
        stmt = stmt.offset(offset).limit(per_page)

        # Execute query
        recordings = db.session.execute(stmt).scalars().all()

        # Enrich recordings with sharing metadata
        enriched_recordings = []
        for recording in recordings:
            rec_dict = recording.to_list_dict(viewer_user=current_user)

            # Add sharing metadata
            is_owner = recording.user_id == current_user.id
            rec_dict['is_owner'] = is_owner

            # Get per-user status (owner uses Recording fields, recipients use SharedRecordingState)
            user_inbox, user_highlighted = get_user_recording_status(recording, current_user)
            rec_dict['is_inbox'] = user_inbox
            rec_dict['is_highlighted'] = user_highlighted

            # Add edit permission info (uses has_recording_access which checks group admin status)
            rec_dict['can_edit'] = has_recording_access(recording, current_user, require_edit=True)

            # Add delete permission info (only owner can delete)
            rec_dict['can_delete'] = is_owner and (USERS_CAN_DELETE or current_user.is_admin)

            if not is_owner:
                # This is a shared recording - get owner info and share permissions
                owner = User.query.get(recording.user_id)
                rec_dict['owner_username'] = owner.username if owner else "Unknown"
                rec_dict['is_shared'] = True
                # Don't show outgoing share count for recordings you don't own
                rec_dict['shared_with_count'] = 0
                rec_dict['public_share_count'] = 0

                # Get share permissions
                share = InternalShare.query.filter_by(
                    recording_id=recording.id,
                    shared_with_user_id=current_user.id
                ).first()

                if share:
                    rec_dict['share_info'] = {
                        'share_id': share.id,
                        'owner_username': owner.username if owner else "Unknown",
                        'can_edit': share.can_edit,
                        'can_reshare': share.can_reshare,
                        'shared_at': share.created_at.isoformat()
                    }
                else:
                    # Fallback if share record not found (shouldn't happen)
                    rec_dict['share_info'] = {
                        'can_edit': False,
                        'can_reshare': False
                    }
            else:
                rec_dict['is_shared'] = False

            # Check if recording has group tags (among visible tags)
            visible_tags = recording.get_visible_tags(current_user)
            has_group_tags = any(tag.is_group_tag for tag in visible_tags)
            rec_dict['has_group_tags'] = has_group_tags

            enriched_recordings.append(rec_dict)

        # Calculate pagination metadata
        total_pages = (total_count + per_page - 1) // per_page
        has_next = page < total_pages
        has_prev = page > 1

        return jsonify({
            'recordings': enriched_recordings,
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total': total_count,
                'total_pages': total_pages,
                'has_next': has_next,
                'has_prev': has_prev
            }
        })

    except Exception as e:
        current_app.logger.error(f"Error fetching paginated recordings: {e}")
        return jsonify({'error': str(e)}), 500


@recordings_bp.route('/save', methods=['POST'])
@login_required
def save_metadata():
    """Save recording metadata (title, participants, notes, summary, etc.)."""
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        recording_id = data.get('id')
        if not recording_id:
            return jsonify({'error': 'No recording ID provided'}), 400

        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Check if user has at least view access
        if not has_recording_access(recording, current_user, require_edit=False):
            return jsonify({'error': 'You do not have permission to access this recording'}), 403

        # Handle notes separately - no edit permission required (user-specific)
        if 'notes' in data:
            if recording.user_id == current_user.id:
                # Owner saves to Recording.notes
                recording.notes = sanitize_html(data['notes']) if data['notes'] else data['notes']
            else:
                # Shared user saves to personal_notes (requires SharedRecordingState)
                from src.models.sharing import SharedRecordingState
                state = SharedRecordingState.query.filter_by(
                    recording_id=recording.id,
                    user_id=current_user.id
                ).first()

                if not state:
                    # Create SharedRecordingState if it doesn't exist
                    state = SharedRecordingState(
                        recording_id=recording.id,
                        user_id=current_user.id,
                        is_inbox=True,
                        is_highlighted=False
                    )
                    db.session.add(state)

                state.personal_notes = sanitize_html(data['notes']) if data['notes'] else data['notes']

        # Determine if any fields requiring edit permission are being updated
        edit_fields = ['title', 'participants', 'summary', 'meeting_date']
        requires_edit = any(field in data for field in edit_fields)

        # If edit fields are present, check for edit permission
        if requires_edit and not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to edit this recording'}), 403

        # Update fields requiring edit permission
        if requires_edit:
            if 'title' in data:
                recording.title = data['title']
            if 'participants' in data:
                recording.participants = data['participants']
            if 'summary' in data:
                recording.summary = sanitize_html(data['summary']) if data['summary'] else data['summary']
            if 'meeting_date' in data:
                try:
                    date_str = data['meeting_date']
                    if date_str:
                        # Try to parse as full ISO datetime first
                        try:
                            recording.meeting_date = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
                        except (ValueError, AttributeError):
                            # Fall back to date-only format, preserve existing time if available
                            parsed_date = datetime.strptime(date_str, '%Y-%m-%d')
                            if recording.meeting_date:
                                # Preserve existing time
                                existing_time = recording.meeting_date.time()
                                recording.meeting_date = datetime.combine(parsed_date.date(), existing_time)
                            else:
                                # No existing time, use the parsed date with midnight time
                                recording.meeting_date = parsed_date
                    else:
                        recording.meeting_date = None
                except (ValueError, TypeError) as e:
                    current_app.logger.warning(f"Could not parse meeting_date '{data.get('meeting_date')}': {e}")

        # Handle per-user status fields (only requires view permission)
        if 'is_inbox' in data or 'is_highlighted' in data:
            set_user_recording_status(
                recording,
                current_user,
                is_inbox=data.get('is_inbox'),
                is_highlighted=data.get('is_highlighted')
            )

        db.session.commit()

        # Re-export the recording if auto-export is enabled and editable fields were changed
        if requires_edit:
            export_recording(recording_id)

        # Return recording with per-user status
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({'success': True, 'recording': recording_dict})

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error saving metadata for recording {data.get('id')}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred while saving.'}), 500


@recordings_bp.route('/recording/<int:recording_id>/update_transcription', methods=['POST'])
@login_required
def update_transcription(recording_id):
    """Updates the transcription content for a recording."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user, require_edit=True):
            return jsonify({'error': 'You do not have permission to edit this recording'}), 403

        data = request.json
        new_transcription = data.get('transcription')

        if new_transcription is None:
            return jsonify({'error': 'No transcription data provided'}), 400

        # The incoming data could be a JSON string (from ASR edit) or plain text
        recording.transcription = new_transcription
        
        # Optional: If the transcription changes, we might want to indicate that the summary is outdated.
        # For now, we'll just save the transcript. A "regenerate summary" button could be a good follow-up.

        db.session.commit()
        current_app.logger.info(f"Transcription for recording {recording_id} was updated.")

        # Re-export the recording if auto-export is enabled
        export_recording(recording_id)

        # Return recording with per-user status
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({'success': True, 'message': 'Transcription updated successfully.', 'recording': recording_dict})

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error updating transcription for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred while updating the transcription.'}), 500

# Toggle inbox status endpoint


@recordings_bp.route('/recording/<int:recording_id>/toggle_inbox', methods=['POST'])
@login_required
def toggle_inbox(recording_id):
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Only require view access (not edit) - users can manage their own inbox status
        if not has_recording_access(recording, current_user, require_edit=False):
            return jsonify({'error': 'You do not have permission to view this recording'}), 403

        # Get current status and toggle it
        current_inbox, current_highlighted = get_user_recording_status(recording, current_user)
        new_inbox, new_highlighted = set_user_recording_status(recording, current_user, is_inbox=not current_inbox)

        return jsonify({'success': True, 'is_inbox': new_inbox})
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error toggling inbox status for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500

# Toggle highlighted status endpoint


@recordings_bp.route('/recording/<int:recording_id>/toggle_highlight', methods=['POST'])
@login_required
def toggle_highlight(recording_id):
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Only require view access (not edit) - users can manage their own highlight status
        if not has_recording_access(recording, current_user, require_edit=False):
            return jsonify({'error': 'You do not have permission to view this recording'}), 403

        # Get current status and toggle it
        current_inbox, current_highlighted = get_user_recording_status(recording, current_user)
        new_inbox, new_highlighted = set_user_recording_status(recording, current_user, is_highlighted=not current_highlighted)

        return jsonify({'success': True, 'is_highlighted': new_highlighted})
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error toggling highlighted status for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500




@recordings_bp.route('/upload', methods=['POST'])
@login_required
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        original_filename = file.filename
        safe_filename = secure_filename(original_filename)
        filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{safe_filename}")

        # Get original file size
        file.seek(0, os.SEEK_END)
        original_file_size = file.tell()
        file.seek(0)

        # Check size limit before saving - only enforce if chunking is disabled or using ASR endpoint
        max_content_length = current_app.config.get('MAX_CONTENT_LENGTH')
        
        # Skip size check if chunking is enabled and using OpenAI Whisper API
        should_enforce_size_limit = True
        if ENABLE_CHUNKING and chunking_service and not USE_ASR_ENDPOINT:
            should_enforce_size_limit = False
            # Get chunking mode for better logging
            mode, limit_value = chunking_service.parse_chunk_limit()
            if mode == 'size':
                current_app.logger.info(f"Size-based chunking enabled ({limit_value}MB limit) - skipping {original_file_size/1024/1024:.1f}MB size limit check")
            else:
                current_app.logger.info(f"Duration-based chunking enabled ({limit_value}s limit) - skipping {original_file_size/1024/1024:.1f}MB size limit check")
        
        if should_enforce_size_limit and max_content_length and original_file_size > max_content_length:
            raise RequestEntityTooLarge()

        file.save(filepath)
        current_app.logger.info(f"File saved to {filepath}")

        # --- Convert files only when chunking is needed ---
        filename_lower = original_filename.lower()
        
        # Check if chunking will be needed for this file
        needs_chunking_for_processing = (chunking_service and 
                                       ENABLE_CHUNKING and 
                                       not USE_ASR_ENDPOINT and
                                       chunking_service.needs_chunking(filepath, USE_ASR_ENDPOINT))
        
        # Define supported formats based on whether chunking is needed
        if needs_chunking_for_processing:
            # For chunking: only support formats that work well with chunking
            supported_formats = ('.wav', '.mp3', '.flac')
            convertible_formats = ('.amr', '.3gp', '.3gpp', '.m4a', '.aac', '.ogg', '.wma', '.webm', '.mp4', '.mov', '.opus', '.caf', '.aiff', '.ts', '.mts', '.mkv', '.avi', '.m4v', '.wmv', '.flv', '.mpeg', '.mpg', '.ogv', '.vob', '.asf')
        else:
            # For direct transcription: support WebM and other formats directly
            supported_formats = ('.wav', '.mp3', '.flac', '.webm', '.m4a', '.aac', '.ogg')
            convertible_formats = ('.amr', '.3gp', '.3gpp', '.wma', '.mp4', '.mov', '.opus', '.caf', '.aiff', '.ts', '.mts', '.mkv', '.avi', '.m4v', '.wmv', '.flv', '.mpeg', '.mpg', '.ogv', '.vob', '.asf')
        
        # Special handling for problematic AAC files when using ASR endpoint
        is_problematic_aac = (USE_ASR_ENDPOINT and 
                             (filename_lower.endswith('.aac') or 
                              'aac' in filename_lower.lower()))
        
        # Convert if file is not in supported formats OR is problematic AAC for ASR
        should_convert = ((not filename_lower.endswith(supported_formats) and needs_chunking_for_processing) or 
                         is_problematic_aac)
        
        if should_convert:
            if is_problematic_aac:
                current_app.logger.info(f"Converting AAC-encoded file {filename_lower} to high-quality MP3 for ASR endpoint compatibility.")
            elif filename_lower.endswith(convertible_formats):
                current_app.logger.info(f"Converting {filename_lower} format to high-quality MP3 for chunking processing.")
            else:
                current_app.logger.info(f"Attempting to convert unknown format ({filename_lower}) to high-quality MP3 for chunking.")
            
            base_filepath, _ = os.path.splitext(filepath)
            temp_mp3_filepath = f"{base_filepath}_temp.mp3"
            mp3_filepath = f"{base_filepath}.mp3"

            try:
                # Convert to high-quality MP3 (128kbps, 44.1kHz) for better transcription accuracy
                subprocess.run(
                    ['ffmpeg', '-i', filepath, '-y', '-acodec', 'libmp3lame', '-b:a', '128k', '-ar', '44100', '-ac', '1', temp_mp3_filepath],
                    check=True, capture_output=True, text=True
                )
                current_app.logger.info(f"Successfully converted {filepath} to {temp_mp3_filepath} (128kbps MP3)")
                
                # If the original file is not the same as the final mp3 file, remove it
                if filepath.lower() != mp3_filepath.lower():
                    os.remove(filepath)
                
                # Rename the temporary file to the final filename
                os.rename(temp_mp3_filepath, mp3_filepath)
                
                filepath = mp3_filepath
            except FileNotFoundError:
                current_app.logger.error("ffmpeg command not found. Please ensure ffmpeg is installed and in the system's PATH.")
                return jsonify({'error': 'Audio conversion tool (ffmpeg) not found on server.'}), 500
            except subprocess.CalledProcessError as e:
                current_app.logger.error(f"ffmpeg conversion failed for {filepath}: {e.stderr}")
                return jsonify({'error': f'Failed to convert audio file: {e.stderr}'}), 500
        elif not filename_lower.endswith(supported_formats):
            # File is not supported and chunking is not needed - log but don't convert
            current_app.logger.info(f"File format {filename_lower} will be processed directly without conversion (chunking not needed)")

        # Get final file size (of original or converted file)
        final_file_size = os.path.getsize(filepath)

        # Determine MIME type of the final file
        mime_type, _ = mimetypes.guess_type(filepath)
        current_app.logger.info(f"Final MIME type: {mime_type} for file {filepath}")

        # Get notes from the form
        notes = request.form.get('notes')
        
        # Get selected tags if provided (multiple tags support)
        selected_tags = []
        tag_index = 0
        while True:
            tag_id_key = f'tag_ids[{tag_index}]'
            tag_id = request.form.get(tag_id_key)
            if not tag_id:
                break

            # Check if tag belongs to user OR is a group tag where user is a member
            tag = Tag.query.filter_by(id=tag_id).first()
            if tag:
                # Allow tag if it's user's own tag OR it's a group tag where user is a member
                if tag.user_id == current_user.id or (tag.group_id and GroupMembership.query.filter_by(group_id=tag.group_id, user_id=current_user.id).first()):
                    selected_tags.append(tag)
            tag_index += 1

        # For backward compatibility with single tag uploads
        if not selected_tags:
            single_tag_id = request.form.get('tag_id')
            if single_tag_id:
                # Check if tag belongs to user OR is a group tag where user is a member
                tag = Tag.query.filter_by(id=single_tag_id).first()
                if tag and (tag.user_id == current_user.id or (tag.group_id and GroupMembership.query.filter_by(group_id=tag.group_id, user_id=current_user.id).first())):
                    selected_tags.append(tag)
        
        # Get ASR advanced options if provided
        language = request.form.get('language', '')
        min_speakers = request.form.get('min_speakers') or None
        max_speakers = request.form.get('max_speakers') or None
        
        # Convert to int if provided
        if min_speakers:
            try:
                min_speakers = int(min_speakers)
            except (ValueError, TypeError):
                min_speakers = None
        if max_speakers:
            try:
                max_speakers = int(max_speakers)
            except (ValueError, TypeError):
                max_speakers = None
        
        # Apply precedence hierarchy: user input > tag defaults > environment variables > auto-detect
        
        # Apply tag defaults if tags are selected and values are not explicitly provided by user
        # Use first tag's defaults (highest priority)
        if selected_tags:
            first_tag = selected_tags[0]
            if not language and first_tag.default_language:
                language = first_tag.default_language
            if min_speakers is None and first_tag.default_min_speakers:
                min_speakers = first_tag.default_min_speakers
            if max_speakers is None and first_tag.default_max_speakers:
                max_speakers = first_tag.default_max_speakers
        
        # Apply environment variable defaults if still no values are set
        if min_speakers is None and ASR_MIN_SPEAKERS:
            try:
                min_speakers = int(ASR_MIN_SPEAKERS)
            except (ValueError, TypeError):
                min_speakers = None
        if max_speakers is None and ASR_MAX_SPEAKERS:
            try:
                max_speakers = int(ASR_MAX_SPEAKERS)
            except (ValueError, TypeError):
                max_speakers = None

        # Create initial database entry
        now = datetime.utcnow()
        recording = Recording(
            audio_path=filepath,
            original_filename=original_filename,
            title=f"Recording - {original_filename}",
            file_size=final_file_size,
            status='PENDING',
            meeting_date=now,
            user_id=current_user.id,
            mime_type=mime_type,
            notes=notes,
            processing_source='upload'  # Track that this was manually uploaded
        )
        db.session.add(recording)
        db.session.commit()
        
        # Add tags to recording if selected (preserve order)
        for order, tag in enumerate(selected_tags, 1):
            new_association = RecordingTag(
                recording_id=recording.id,
                tag_id=tag.id,
                order=order,
                added_at=datetime.utcnow()
            )
            db.session.add(new_association)
        
        if selected_tags:
            db.session.commit()
            tag_names = [tag.name for tag in selected_tags]
            current_app.logger.info(f"Added {len(selected_tags)} tags to recording {recording.id}: {', '.join(tag_names)}")
        
        current_app.logger.info(f"Initial recording record created with ID: {recording.id}")

        # --- Queue transcription job ---
        first_tag = selected_tags[0] if selected_tags else None
        job_params = {
            'language': language,
            'min_speakers': min_speakers,
            'max_speakers': max_speakers,
            'tag_id': first_tag.id if first_tag else None
        }

        current_app.logger.info(f"Queueing transcription for recording {recording.id} with params: {job_params}")
        job_queue.enqueue(
            user_id=current_user.id,
            recording_id=recording.id,
            job_type='transcribe',
            params=job_params,
            is_new_upload=True
        )
        current_app.logger.info(f"Transcription job queued for recording ID: {recording.id}")

        return jsonify(recording.to_dict(viewer_user=current_user)), 202

    except RequestEntityTooLarge:
        max_size_mb = current_app.config['MAX_CONTENT_LENGTH'] / (1024 * 1024)
        current_app.logger.warning(f"Upload failed: File too large (>{max_size_mb}MB)")
        return jsonify({
            'error': f'File too large. Maximum size is {max_size_mb:.0f} MB.',
            'max_size_mb': max_size_mb
        }), 413
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error during file upload: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred during upload.'}), 500


# Status Endpoint


@recordings_bp.route('/recording/<int:recording_id>', methods=['DELETE'])
@login_required
def delete_recording(recording_id):
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404
            
        # Check if the recording belongs to the current user
        if recording.user_id and recording.user_id != current_user.id:
            return jsonify({'error': 'You do not have permission to delete this recording'}), 403

        # Check deletion permissions (admin-only if USERS_CAN_DELETE is false)
        if not USERS_CAN_DELETE and not current_user.is_admin:
            return jsonify({'error': 'Only administrators can delete recordings'}), 403

        # Delete the audio file first
        try:
            if recording.audio_path and os.path.exists(recording.audio_path):
                os.remove(recording.audio_path)
                current_app.logger.info(f"Deleted audio file: {recording.audio_path}")
        except Exception as e:
            current_app.logger.error(f"Error deleting audio file {recording.audio_path}: {e}")

        # Log embeddings cleanup for Inquire Mode if enabled
        if ENABLE_INQUIRE_MODE:
            chunk_count = TranscriptChunk.query.filter_by(recording_id=recording_id).count()
            if chunk_count > 0:
                current_app.logger.info(f"Deleting {chunk_count} transcript chunks with embeddings for recording {recording_id}")

        # Delete associated processing jobs (required because recording_id is NOT NULL)
        from src.models.processing_job import ProcessingJob
        deleted_jobs = ProcessingJob.query.filter_by(recording_id=recording_id).delete()
        if deleted_jobs > 0:
            current_app.logger.info(f"Deleted {deleted_jobs} processing jobs for recording {recording_id}")

        # Delete the database record (cascade will handle chunks/embeddings)
        db.session.delete(recording)
        db.session.commit()
        current_app.logger.info(f"Deleted recording record ID: {recording_id}")

        if ENABLE_INQUIRE_MODE and chunk_count > 0:
            current_app.logger.info(f"Successfully deleted embeddings and chunks for recording {recording_id}")

        # Mark the export file as deleted
        mark_export_as_deleted(recording_id)

        # Clean up orphaned speakers (run after successful deletion)
        # This is a best-effort cleanup; failures are logged but don't affect the delete operation
        try:
            from src.services.speaker_cleanup import cleanup_orphaned_speakers
            speaker_stats = cleanup_orphaned_speakers()
            if speaker_stats.get('speakers_deleted', 0) > 0:
                current_app.logger.info(
                    f"Cleaned up {speaker_stats['speakers_deleted']} orphaned speakers after recording deletion"
                )
        except Exception as cleanup_error:
            # Log the error but don't fail the deletion
            current_app.logger.warning(f"Speaker cleanup after recording deletion failed: {cleanup_error}")

        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred while deleting.'}), 500


# --- Inbox and Archive Endpoints ---

@recordings_bp.route('/api/inbox_recordings', methods=['GET'])
@login_required
def get_inbox_recordings():
    """Get recordings that are in the inbox and currently processing."""
    from sqlalchemy import select
    try:
        stmt = select(Recording).where(
            Recording.user_id == current_user.id,
            Recording.is_inbox == True,
            Recording.status.in_(['PENDING', 'PROCESSING', 'SUMMARIZING'])
        ).order_by(Recording.created_at.desc())

        recordings = db.session.execute(stmt).scalars().all()
        return jsonify([recording.to_list_dict(viewer_user=current_user) for recording in recordings])
    except Exception as e:
        current_app.logger.error(f"Error fetching inbox recordings: {e}")
        return jsonify({'error': str(e)}), 500


@recordings_bp.route('/api/recordings/archived', methods=['GET'])
@login_required
def get_archived_recordings():
    """Get recordings where audio has been deleted but transcription remains."""
    from sqlalchemy import select
    try:
        search_query = request.args.get('q', '').strip()

        # Find recordings owned by current user where audio_deleted_at is not null
        stmt = select(Recording).where(
            Recording.user_id == current_user.id,
            Recording.audio_deleted_at.is_not(None)
        ).order_by(Recording.audio_deleted_at.desc())

        recordings = db.session.execute(stmt).scalars().all()
        return jsonify([recording.to_list_dict(viewer_user=current_user) for recording in recordings])
    except Exception as e:
        current_app.logger.error(f"Error fetching archived recordings: {e}")
        return jsonify({'error': str(e)}), 500


# --- Recording Detail and Audio Endpoints ---

@recordings_bp.route('/api/recordings/<int:recording_id>', methods=['GET'])
@login_required
def get_recording_detail(recording_id):
    """Get full details for a specific recording including markdown HTML."""
    try:
        recording = db.session.get(Recording, recording_id)

        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Check ownership or shared access
        has_access = recording.user_id == current_user.id

        # Check if recording has been shared with current user (if internal sharing is enabled)
        if not has_access and ENABLE_INTERNAL_SHARING:
            share = InternalShare.query.filter_by(
                recording_id=recording_id,
                shared_with_user_id=current_user.id
            ).first()
            has_access = share is not None

        if not has_access:
            return jsonify({'error': 'Access denied'}), 403

        # Return full detail with HTML conversion
        rec_dict = recording.to_dict(include_html=True, viewer_user=current_user)

        # Add sharing metadata
        is_owner = recording.user_id == current_user.id
        rec_dict['is_owner'] = is_owner

        # Add edit permission info (uses has_recording_access which checks group admin status)
        rec_dict['can_edit'] = has_recording_access(recording, current_user, require_edit=True)

        # Add delete permission info (only owner can delete)
        rec_dict['can_delete'] = is_owner and (USERS_CAN_DELETE or current_user.is_admin)

        # Add sharing-related fields
        if not is_owner:
            # This is a shared recording - get owner info and share permissions
            owner = User.query.get(recording.user_id)
            rec_dict['owner_username'] = owner.username if owner else "Unknown"
            rec_dict['is_shared'] = True
            # Don't show outgoing share count for recordings you don't own
            rec_dict['shared_with_count'] = 0
            rec_dict['public_share_count'] = 0

            # Get share permissions
            share = InternalShare.query.filter_by(
                recording_id=recording.id,
                shared_with_user_id=current_user.id
            ).first()

            if share:
                rec_dict['share_info'] = {
                    'share_id': share.id,
                    'owner_username': owner.username if owner else "Unknown",
                    'can_edit': share.can_edit,
                    'can_reshare': share.can_reshare,
                    'shared_at': share.created_at.isoformat()
                }
            else:
                # Fallback if share record not found (shouldn't happen)
                rec_dict['share_info'] = {
                    'can_edit': False,
                    'can_reshare': False
                }
        else:
            rec_dict['is_shared'] = False

        # Check if recording has group tags (among visible tags)
        visible_tags = recording.get_visible_tags(current_user)
        has_group_tags = any(tag.is_group_tag for tag in visible_tags) if visible_tags else False
        rec_dict['has_group_tags'] = has_group_tags

        # Enrich with per-user status
        enrich_recording_dict_with_user_status(rec_dict, recording, current_user)

        return jsonify(rec_dict)
    except Exception as e:
        current_app.logger.error(f"Error fetching recording detail: {e}")
        return jsonify({'error': str(e)}), 500


@recordings_bp.route('/recording/<int:recording_id>/status', methods=['GET'])
@login_required
def get_recording_status_only(recording_id):
    """
    Lightweight endpoint that returns only the status field.
    Used for polling during processing/summarization.
    Note: Rate limiting exemption is configured at app level.
    """
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user):
            return jsonify({'error': 'You do not have permission to view this recording'}), 403

        # Return only the status field
        return jsonify({'status': recording.status})
    except Exception as e:
        current_app.logger.error(f"Error fetching status for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/api/recordings/batch-status', methods=['POST'])
@login_required
def get_batch_recording_status():
    """
    Batch endpoint to get status for multiple recordings at once.
    More efficient than polling individual status endpoints.

    Request body: {"recording_ids": [1, 2, 3]}
    Response: {"statuses": {"1": "COMPLETED", "2": "PROCESSING", "3": "FAILED"}}
    """
    try:
        data = request.get_json()
        if not data or 'recording_ids' not in data:
            return jsonify({'error': 'recording_ids is required'}), 400

        recording_ids = data['recording_ids']
        if not isinstance(recording_ids, list):
            return jsonify({'error': 'recording_ids must be a list'}), 400

        # Limit batch size to prevent abuse
        if len(recording_ids) > 50:
            return jsonify({'error': 'Maximum 50 recordings per batch'}), 400

        # Query all recordings at once
        recordings = Recording.query.filter(Recording.id.in_(recording_ids)).all()

        # Build response with only accessible recordings
        statuses = {}
        for recording in recordings:
            if has_recording_access(recording, current_user):
                statuses[str(recording.id)] = recording.status

        return jsonify({'statuses': statuses})
    except Exception as e:
        current_app.logger.error(f"Error fetching batch status: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/api/recordings/job-queue-status', methods=['GET'])
@login_required
def get_job_queue_status():
    """
    Get detailed job queue status for all jobs (active, completed, and failed).
    Returns status for the user's jobs within the last hour.
    """
    try:
        from src.models import ProcessingJob
        from src.services.job_queue import TRANSCRIPTION_JOBS, SUMMARY_JOBS
        from datetime import timedelta

        # Expire all cached objects to ensure we see latest data from worker threads
        db.session.expire_all()

        # Get all jobs for the user (active + recent completed/failed within last hour)
        cutoff_time = datetime.utcnow() - timedelta(hours=1)
        all_jobs = ProcessingJob.query.filter(
            ProcessingJob.user_id == current_user.id,
            db.or_(
                ProcessingJob.status.in_(['queued', 'processing']),
                db.and_(
                    ProcessingJob.status.in_(['completed', 'failed']),
                    ProcessingJob.completed_at >= cutoff_time
                )
            )
        ).order_by(ProcessingJob.created_at.desc()).all()

        job_details = []
        for job in all_jobs:
            recording = db.session.get(Recording, job.recording_id)
            recording_title = None
            if recording:
                recording_title = recording.title or recording.original_filename or 'Untitled'

            # Determine queue type
            queue_type = 'summary' if job.job_type in SUMMARY_JOBS else 'transcription'

            # Calculate position if queued
            position = None
            if job.status == 'queued':
                job_types = SUMMARY_JOBS if job.job_type in SUMMARY_JOBS else TRANSCRIPTION_JOBS
                ahead_in_queue = ProcessingJob.query.filter(
                    ProcessingJob.status == 'queued',
                    ProcessingJob.job_type.in_(job_types),
                    ProcessingJob.created_at < job.created_at
                ).count()
                currently_processing = ProcessingJob.query.filter(
                    ProcessingJob.status == 'processing',
                    ProcessingJob.job_type.in_(job_types)
                ).count()
                position = ahead_in_queue + currently_processing + 1

            job_details.append({
                'id': job.id,
                'recording_id': job.recording_id,
                'recording_title': recording_title,
                'job_status': job.status,
                'job_type': job.job_type,
                'queue_type': queue_type,
                'position': position,
                'is_new_upload': job.is_new_upload,
                'error_message': job.error_message,
                'created_at': job.created_at.isoformat() if job.created_at else None,
                'started_at': job.started_at.isoformat() if job.started_at else None,
                'completed_at': job.completed_at.isoformat() if job.completed_at else None
            })

        return jsonify({'jobs': job_details})
    except Exception as e:
        current_app.logger.error(f"Error fetching job queue status: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/api/recordings/jobs/<int:job_id>/retry', methods=['POST'])
@login_required
def retry_failed_job(job_id):
    """Retry a failed job."""
    try:
        from src.models import ProcessingJob

        job = db.session.get(ProcessingJob, job_id)
        if not job:
            return jsonify({'error': 'Job not found'}), 404

        if job.user_id != current_user.id:
            return jsonify({'error': 'Access denied'}), 403

        if job.status != 'failed':
            return jsonify({'error': 'Only failed jobs can be retried'}), 400

        # Reset job for retry
        job.status = 'queued'
        job.error_message = None
        job.retry_count = 0
        job.started_at = None
        job.completed_at = None
        db.session.commit()

        current_app.logger.info(f"Job {job_id} queued for retry by user {current_user.id}")
        return jsonify({'success': True, 'message': 'Job queued for retry'})

    except Exception as e:
        current_app.logger.error(f"Error retrying job {job_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/api/recordings/jobs/<int:job_id>', methods=['DELETE'])
@login_required
def delete_job(job_id):
    """Delete a job (clear from queue or history)."""
    try:
        from src.models import ProcessingJob
        import os

        job = db.session.get(ProcessingJob, job_id)
        if not job:
            return jsonify({'error': 'Job not found'}), 404

        if job.user_id != current_user.id:
            return jsonify({'error': 'Access denied'}), 403

        # If it's a failed new upload, also delete the recording
        if job.status == 'failed' and job.is_new_upload:
            recording = db.session.get(Recording, job.recording_id)
            if recording:
                # Delete audio file
                if recording.audio_path and os.path.exists(recording.audio_path):
                    try:
                        os.remove(recording.audio_path)
                    except Exception as e:
                        current_app.logger.error(f"Error deleting audio file: {e}")
                # Delete ALL processing jobs for this recording first
                ProcessingJob.query.filter_by(recording_id=recording.id).delete()
                db.session.delete(recording)
        else:
            # Just delete this job
            db.session.delete(job)
        db.session.commit()

        current_app.logger.info(f"Job {job_id} deleted by user {current_user.id}")
        return jsonify({'success': True, 'message': 'Job deleted'})

    except Exception as e:
        current_app.logger.error(f"Error deleting job {job_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/api/recordings/jobs/clear-completed', methods=['POST'])
@login_required
def clear_completed_jobs():
    """Clear all completed jobs for the current user."""
    try:
        from src.models import ProcessingJob

        deleted = ProcessingJob.query.filter(
            ProcessingJob.user_id == current_user.id,
            ProcessingJob.status == 'completed'
        ).delete(synchronize_session=False)

        db.session.commit()
        current_app.logger.info(f"Cleared {deleted} completed jobs for user {current_user.id}")
        return jsonify({'success': True, 'deleted': deleted})

    except Exception as e:
        current_app.logger.error(f"Error clearing completed jobs: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/status/<int:recording_id>', methods=['GET'])
@login_required
def get_status(recording_id):
    """Endpoint to check the transcription/summarization status (full recording data)."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user):
            return jsonify({'error': 'You do not have permission to view this recording'}), 403

        # Ensure events are loaded (refresh the recording to get latest relationships)
        db.session.refresh(recording)

        # Get recording dict and enrich with per-user status
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)

        return jsonify(recording_dict)
    except Exception as e:
        current_app.logger.error(f"Error fetching status for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/audio/<int:recording_id>')
@login_required
def get_audio(recording_id):
    """Serve audio file for a recording."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording or not recording.audio_path:
            return jsonify({'error': 'Recording or audio file not found'}), 404

        # Check if the recording belongs to the current user or has been shared with them
        has_access = recording.user_id == current_user.id

        # Check if recording has been shared with current user (if internal sharing is enabled)
        if not has_access and ENABLE_INTERNAL_SHARING:
            share = InternalShare.query.filter_by(
                recording_id=recording_id,
                shared_with_user_id=current_user.id
            ).first()
            has_access = share is not None

        if not has_access:
            return jsonify({'error': 'You do not have permission to access this audio file'}), 403
        if not os.path.exists(recording.audio_path):
            current_app.logger.error(f"Audio file missing from server: {recording.audio_path}")
            return jsonify({'error': 'Audio file missing from server'}), 404
        return send_file(recording.audio_path)
    except Exception as e:
        current_app.logger.error(f"Error serving audio for recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


# --- Chat with Transcription ---

@recordings_bp.route('/chat', methods=['POST'])
@login_required
def chat_with_transcription():
    """Chat with a specific recording's transcription."""
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        recording_id = data.get('recording_id')
        user_message = data.get('message')
        message_history = data.get('message_history', [])

        if not recording_id:
            return jsonify({'error': 'No recording ID provided'}), 400
        if not user_message:
            return jsonify({'error': 'No message provided'}), 400

        # Get the recording
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if not has_recording_access(recording, current_user):
            return jsonify({'error': 'You do not have permission to chat with this recording'}), 403

        # Check if chat client is available
        if chat_client is None:
            return jsonify({'error': 'Chat service is not available (chat client not configured)'}), 503

        # Prepare the system prompt with the transcription
        user_chat_output_language = current_user.output_language if current_user.is_authenticated else None

        language_instruction = ""
        if user_chat_output_language:
            language_instruction = f"Please provide all your responses in {user_chat_output_language}."

        user_name = current_user.name if current_user.is_authenticated and current_user.name else "User"
        user_title = current_user.job_title if current_user.is_authenticated and current_user.job_title else "a professional"
        user_company = current_user.company if current_user.is_authenticated and current_user.company else "their organization"

        formatted_transcription = format_transcription_for_llm(recording.transcription)

        # Get configurable transcript length limit for chat
        transcript_limit = SystemSetting.get_setting('transcript_length_limit', 30000)
        if transcript_limit == -1:
            # No limit
            chat_transcript = formatted_transcription
        else:
            chat_transcript = formatted_transcription[:transcript_limit]

        system_prompt = f"""You are a professional meeting and audio transcription analyst assisting {user_name}, who is a(n) {user_title} at {user_company}. {language_instruction} Analyze the following meeting information and respond to the specific request.

Following are the meeting participants and their roles:
{recording.participants or "No specific participants information provided."}

Following is the meeting transcript:
<<start transcript>>
{chat_transcript or "No transcript available."}
<<end transcript>>

Additional context and notes about the meeting:
{recording.notes or "none"}
"""

        # Prepare messages array with system prompt and conversation history
        messages = [{"role": "system", "content": system_prompt}]
        if message_history:
            messages.extend(message_history)
        messages.append({"role": "user", "content": user_message})

        def generate():
            try:
                # Enable streaming
                stream = call_chat_completion(
                    messages=messages,
                    temperature=0.7,
                    max_tokens=int(os.environ.get("CHAT_MAX_TOKENS", "2000")),
                    stream=True
                )

                # Use helper function to process streaming with thinking tag support
                for response in process_streaming_with_thinking(stream):
                    yield response

            except Exception as e:
                current_app.logger.error(f"Error during chat stream generation: {str(e)}")
                # Yield an error message in SSE format
                yield f"data: {json.dumps({'error': str(e)})}\n\n"

        return Response(generate(), mimetype='text/event-stream')

    except Exception as e:
        current_app.logger.error(f"Error in chat endpoint: {str(e)}")
        return jsonify({'error': str(e)}), 500


# --- Tag Management for Recordings ---

@recordings_bp.route('/api/recordings/<int:recording_id>/tags', methods=['POST'])
@login_required
def add_tag_to_recording(recording_id):
    """Add a tag to a recording. Triggers auto-share for group tags."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Check if user has view access to this recording
        # (Edit permission will be checked for group tags specifically)
        if not has_recording_access(recording, current_user, require_edit=False):
            return jsonify({'error': 'You do not have permission to access this recording'}), 403

        data = request.get_json()
        tag_id = data.get('tag_id')

        if not tag_id:
            return jsonify({'error': 'Tag ID is required'}), 400

        tag = db.session.get(Tag, tag_id)
        if not tag:
            return jsonify({'error': 'Tag not found'}), 404

        # Check if user has access to this tag and permission to apply it
        if tag.group_id:
            # Group tag - check membership first
            membership = GroupMembership.query.filter_by(
                group_id=tag.group_id,
                user_id=current_user.id
            ).first()
            if not membership:
                return jsonify({'error': 'You do not have access to this tag'}), 403

            # Only file owner or group admin can apply group tags
            if recording.user_id != current_user.id and membership.role != 'admin':
                return jsonify({'error': 'Only recording owner or group admin can apply group tags'}), 403

            # Group tags require edit permission
            if not has_recording_access(recording, current_user, require_edit=True):
                return jsonify({'error': 'You do not have permission to apply group tags to this recording'}), 403
        else:
            # Personal tag - only the tag owner can use it (view access is sufficient)
            if tag.user_id != current_user.id:
                return jsonify({'error': 'You can only apply your own personal tags'}), 403

        # Check if tag is already on the recording
        existing = RecordingTag.query.filter_by(
            recording_id=recording_id,
            tag_id=tag_id
        ).first()

        if existing:
            return jsonify({'error': 'Tag is already on this recording'}), 400

        # Get the next order position
        max_order = db.session.query(db.func.max(RecordingTag.order)).filter_by(
            recording_id=recording_id
        ).scalar() or 0

        # Add the tag
        recording_tag = RecordingTag(
            recording_id=recording_id,
            tag_id=tag_id,
            order=max_order + 1
        )
        db.session.add(recording_tag)

        # If this is a group tag with sharing enabled, automatically share the recording
        # Only auto-share if recording is completed (not during processing)
        if tag.group_id and ENABLE_INTERNAL_SHARING and recording.status == 'COMPLETED' and (tag.auto_share_on_apply or tag.share_with_group_lead):
            # Determine who to share with
            if tag.auto_share_on_apply:
                group_members = GroupMembership.query.filter_by(group_id=tag.group_id).all()
            elif tag.share_with_group_lead:
                group_members = GroupMembership.query.filter_by(group_id=tag.group_id, role='admin').all()
            else:
                group_members = []

            shares_created = 0
            for membership_to_share in group_members:
                # Skip the recording owner
                if membership_to_share.user_id == recording.user_id:
                    continue

                # Check if already shared
                existing_share = InternalShare.query.filter_by(
                    recording_id=recording_id,
                    shared_with_user_id=membership_to_share.user_id
                ).first()

                if not existing_share:
                    # Create internal share with correct permissions
                    # Group admins get edit permission, regular members get read-only
                    share = InternalShare(
                        recording_id=recording_id,
                        owner_id=recording.user_id,
                        shared_with_user_id=membership_to_share.user_id,
                        can_edit=(membership_to_share.role == 'admin'),
                        can_reshare=False,
                        source_type='group_tag',
                        source_tag_id=tag.id
                    )
                    db.session.add(share)

                    # Check if SharedRecordingState already exists (might exist from previous share)
                    existing_state = SharedRecordingState.query.filter_by(
                        recording_id=recording_id,
                        user_id=membership_to_share.user_id
                    ).first()

                    if not existing_state:
                        # Create SharedRecordingState with default values for the recipient
                        state = SharedRecordingState(
                            recording_id=recording_id,
                            user_id=membership_to_share.user_id,
                            is_inbox=True,  # New shares appear in inbox by default
                            is_highlighted=False  # Not favorited by default
                        )
                        db.session.add(state)

                    shares_created += 1
                    current_app.logger.info(f"Auto-shared recording {recording_id} with user {membership_to_share.user_id} (role={membership_to_share.role}) via group tag '{tag.name}'")

            if shares_created > 0:
                current_app.logger.info(f"Created {shares_created} auto-shares for recording {recording_id} via group tag '{tag.name}'")

        db.session.commit()

        # Return updated recording with per-user status
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({
            'success': True,
            'recording': recording_dict,
            'tag': tag.to_dict()
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error adding tag to recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


@recordings_bp.route('/api/recordings/<int:recording_id>/tags/<int:tag_id>', methods=['DELETE'])
@login_required
def remove_tag_from_recording(recording_id, tag_id):
    """Remove a tag from a recording. Cleans up auto-shares for group tags."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Check if user has view access to this recording
        # (Edit permission will be checked for group tags specifically)
        if not has_recording_access(recording, current_user, require_edit=False):
            return jsonify({'error': 'You do not have permission to access this recording'}), 403

        # Find the recording-tag association
        recording_tag = RecordingTag.query.filter_by(
            recording_id=recording_id,
            tag_id=tag_id
        ).first()

        if not recording_tag:
            return jsonify({'error': 'Tag is not on this recording'}), 404

        # Get the tag to check permissions and for cleanup
        tag = db.session.get(Tag, tag_id)
        if tag:
            # Check permissions to remove this specific tag
            if tag.group_id:
                # Group tag - only file owner or group admin can remove
                membership = GroupMembership.query.filter_by(
                    group_id=tag.group_id,
                    user_id=current_user.id
                ).first()
                if recording.user_id != current_user.id:
                    if not membership or membership.role != 'admin':
                        return jsonify({'error': 'Only recording owner or group admin can remove group tags'}), 403

                # Group tags require edit permission
                if not has_recording_access(recording, current_user, require_edit=True):
                    return jsonify({'error': 'You do not have permission to remove group tags from this recording'}), 403
            else:
                # Personal tag - can be removed by tag owner (view access) or recording owner (edit access)
                if tag.user_id != current_user.id:
                    # Not the tag owner, must be recording owner with edit permission
                    if not has_recording_access(recording, current_user, require_edit=True):
                        return jsonify({'error': 'You can only remove your own personal tags'}), 403

        # Remove the association
        db.session.delete(recording_tag)

        # Clean up shares created by this group tag
        if tag and tag.group_id and ENABLE_INTERNAL_SHARING:
            shares_to_check = InternalShare.query.filter_by(
                recording_id=recording_id,
                source_tag_id=tag_id
            ).all()

            shares_removed = 0
            for share in shares_to_check:
                # Check if user still has access via another group tag on this recording
                other_team_tag_access = db.session.query(Tag).join(
                    RecordingTag, RecordingTag.tag_id == Tag.id
                ).join(
                    GroupMembership, GroupMembership.group_id == Tag.group_id
                ).filter(
                    RecordingTag.recording_id == recording_id,
                    GroupMembership.user_id == share.shared_with_user_id,
                    Tag.id != tag_id,  # Exclude the tag being removed
                    Tag.group_id.isnot(None),
                    db.or_(Tag.auto_share_on_apply == True, Tag.share_with_group_lead == True)
                ).first()

                # Only remove share if user has no other group tag access
                if not other_team_tag_access:
                    db.session.delete(share)
                    shares_removed += 1
                    current_app.logger.info(f"Removed auto-share for user {share.shared_with_user_id} from recording {recording_id} (group tag '{tag.name}' removed)")

            if shares_removed > 0:
                current_app.logger.info(f"Cleaned up {shares_removed} auto-shares for recording {recording_id} after removing group tag '{tag.name}'")

        db.session.commit()

        # Return updated recording with per-user status
        recording_dict = recording.to_dict(viewer_user=current_user)
        enrich_recording_dict_with_user_status(recording_dict, recording, current_user)
        return jsonify({
            'success': True,
            'recording': recording_dict
        })

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error removing tag from recording {recording_id}: {e}", exc_info=True)
        return jsonify({'error': 'An unexpected error occurred.'}), 500


# --- Auto-deletion and Chunks Processing ---

@recordings_bp.route('/api/recordings/<int:recording_id>/toggle_deletion_exempt', methods=['POST'])
@login_required
def toggle_deletion_exempt(recording_id):
    """Toggle the deletion_exempt flag for a recording."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        # Check ownership
        if recording.user_id != current_user.id and not current_user.is_admin:
            return jsonify({'error': 'Permission denied'}), 403

        # Toggle the flag
        recording.deletion_exempt = not recording.deletion_exempt
        db.session.commit()

        return jsonify({
            'success': True,
            'deletion_exempt': recording.deletion_exempt
        })
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error toggling deletion exempt for recording {recording_id}: {e}")
        return jsonify({'error': str(e)}), 500


@recordings_bp.route('/api/recording/<int:recording_id>/process_chunks', methods=['POST'])
@login_required
def process_recording_chunks_endpoint(recording_id):
    """Process chunks for a specific recording."""
    try:
        recording = db.session.get(Recording, recording_id)
        if not recording:
            return jsonify({'error': 'Recording not found'}), 404

        if recording.user_id != current_user.id:
            return jsonify({'error': 'Permission denied'}), 403

        success = process_recording_chunks(recording_id)
        if success:
            return jsonify({'message': 'Chunks processed successfully'})
        else:
            return jsonify({'error': 'Failed to process chunks'}), 500

    except Exception as e:
        current_app.logger.error(f"Error in process chunks endpoint: {e}")
        return jsonify({'error': str(e)}), 500


# --- Inquire Mode API Endpoints ---



